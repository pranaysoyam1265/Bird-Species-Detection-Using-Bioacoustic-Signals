"""
generate_report.py
Generates Professional Technical Report in DOCX and PDF formats

Creates a complete technical report for the Bird Species Detection project
with proper formatting, tables, and figure placeholders.

Author: Bird Detection Project
Date: February 2026
"""

import os
import json
from pathlib import Path
from datetime import datetime
from io import BytesIO

# DOCX generation
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# PDF generation
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4, letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch, cm
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT, TA_RIGHT
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Image, Table, TableStyle,
    PageBreak, ListFlowable, ListItem, KeepTogether
)
from reportlab.platypus.tableofcontents import TableOfContents

import matplotlib.pyplot as plt
import matplotlib
matplotlib.use('Agg')  # Non-interactive backend

# ============================================================
# CONFIGURATION
# ============================================================

BASE_DIR = Path(r"C:\Users\prana\OneDrive\Desktop\ML Conf-BioFSL")
OUTPUT_DIR = BASE_DIR / "10_Outputs" / "Reports"

# Report metadata
REPORT_TITLE = "Confidence-Aware, Explainable Multi-Species Bird Presence Detection System Using Bioacoustic Signals"
AUTHOR_NAME = "[Your Name]"
INSTITUTION = "[Your Institution]"
DATE = datetime.now().strftime("%B %Y")

# Project metrics (update these with your actual values)
METRICS = {
    "total_recordings": 4521,
    "usable_recordings": 1926,
    "num_species": 54,
    "total_chunks": 34811,
    "train_chunks": 23839,
    "val_chunks": 5568,
    "test_chunks": 5404,
    "top1_accuracy": 72.26,
    "top3_accuracy": 82.83,
    "top5_accuracy": 85.55,
    "f1_macro": 63.21,
    "f1_weighted": 72.26,
    "best_epoch": 10,
    "best_val_accuracy": 68.79,
    "train_accuracy": 99.83,
    "training_time": "~2 hours",
    "gpu": "NVIDIA RTX 3050 (4GB)",
}

# ============================================================
# REPORT CONTENT
# ============================================================

ABSTRACT = """
This project presents a deep learning-based system for automatic bird species detection from audio recordings. The system addresses the challenge of manually analyzing long environmental recordings by providing automated species identification with confidence scores and visual explanations.

The system uses mel-spectrogram representation of audio signals and employs an EfficientNet-B0 convolutional neural network for classification. The model was trained on 1,926 recordings from the Xeno-Canto database, covering 54 North American bird species.

Key Results:
• Top-1 Accuracy: 72.26%
• Top-5 Accuracy: 85.55%
• F1 Score (Weighted): 72.26%
• Real-world validation: Correctly identified unseen recordings

The system features Gradient-weighted Class Activation Mapping (Grad-CAM) for explainability, showing which parts of the spectrogram influenced the model's decision. A web-based interface built with Streamlit allows researchers to upload recordings, search for target species, and generate PDF reports.
"""

KEYWORDS = "Bioacoustics, Bird Detection, Deep Learning, EfficientNet, Grad-CAM, Explainable AI, Convolutional Neural Networks, Audio Classification"

INTRODUCTION = {
    "background": """
Bird species monitoring is crucial for ecological research, conservation efforts, and biodiversity assessment. Traditional methods rely on expert ornithologists manually listening to recordings, which presents several challenges:

• Time-consuming: Hours of recordings must be analyzed for minutes of actual bird calls
• Expensive: Requires trained experts with specialized knowledge
• Subjective: Different experts may disagree on species identification
• Limited scalability: Cannot process large datasets efficiently

Automated acoustic monitoring using machine learning offers a solution by processing recordings quickly and consistently. Recent advances in deep learning, particularly convolutional neural networks (CNNs), have shown remarkable success in audio classification tasks when combined with spectrogram representations.
""",
    
    "problem_statement": """
A field researcher collects audio recordings in a forest environment. Each recording typically contains:
• Background noise (wind, insects, rain, human activity)
• Potentially multiple bird species calling simultaneously
• Variable recording quality depending on equipment and conditions

The researcher needs to answer several key questions:
1. Is my target species present in this recording? (Binary: Yes/No)
2. How confident is the detection? (Probability score)
3. Where in the recording does the call occur? (Temporal localization)
4. Why does the system think this? (Explainable evidence)

Current commercial solutions like BirdNET and Merlin provide species identification but lack explainability features that allow researchers to verify and understand the basis for predictions.
""",
    
    "objectives": """
This project aims to build a comprehensive system that:

1. Automatically detects bird species from audio recordings with high accuracy
2. Provides calibrated confidence scores for each detection
3. Localizes detections temporally within the recording
4. Explains predictions through visual heatmaps (Grad-CAM)
5. Offers an intuitive web interface accessible to non-technical researchers
6. Generates professional PDF reports for documentation
7. Provides uncertainty warnings when predictions may be unreliable
""",
    
    "scope": """
The scope of this project includes:
• Dataset: Xeno-Canto recordings (Creative Commons licensed)
• Species coverage: 54 North American bird species
• Input formats: WAV, MP3, OGG, FLAC audio files
• Output: Species identification with confidence scores, temporal localization, and visual explanations
• Deployment: Web-based interface using Streamlit framework
"""
}

METHODOLOGY = {
    "system_architecture": """
The complete system pipeline consists of the following stages:

1. Audio Input: Users upload recordings in various formats (WAV, MP3, OGG, FLAC)
2. Preprocessing: Audio is resampled to 22,050 Hz, converted to mono, and split into 5-second chunks with 50% overlap
3. Feature Extraction: Each chunk is converted to a mel-spectrogram representation
4. Classification: EfficientNet-B0 neural network processes spectrograms and outputs class probabilities
5. Explainability: Grad-CAM generates attention heatmaps showing important regions
6. Aggregation: Chunk-level predictions are aggregated to recording-level results
7. Output: Species predictions with confidence scores, timelines, and visual explanations
""",
    
    "dataset": """
The dataset was sourced from Xeno-Canto (xeno-canto.org), a citizen science project hosting bird sound recordings from around the world under Creative Commons licenses.

Dataset Statistics:
• Total recordings downloaded: 4,521
• Usable recordings (with complete metadata): 1,926
• Species covered: 54 North American bird species
• Total audio duration: Approximately 40 hours
• Geographic focus: North America

Data was split as follows:
• Training set: 70% (1,338 recordings, 23,839 chunks)
• Validation set: 15% (287 recordings, 5,568 chunks)
• Test set: 15% (287 recordings, 5,404 chunks)

The split was performed at the recording level (not chunk level) to prevent data leakage, with stratification by species to ensure balanced representation across splits.
""",
    
    "preprocessing": """
Audio preprocessing involved three main steps:

Step 1: Standardization
• Sample rate: 22,050 Hz (sufficient to capture frequencies up to 11 kHz, covering the typical bird call range of 1-10 kHz)
• Channels: Converted to mono
• Bit depth: 16-bit PCM

Step 2: Chunking
• Duration: 5 seconds per chunk (long enough to capture complete bird calls)
• Overlap: 50% (2.5-second hop) to ensure calls at chunk boundaries are not missed
• Result: 34,811 total chunks from 1,926 recordings

Step 3: Quality Filtering
• Silent chunks removed (RMS amplitude < 0.001)
• Corrupted files identified and excluded
• File integrity verified
""",
    
    "feature_extraction": """
Audio waveforms were converted to mel-spectrograms using the Short-Time Fourier Transform (STFT) with mel-scale frequency mapping.

The mel scale is a perceptual scale that approximates human hearing, providing higher resolution at lower frequencies where bird calls contain important distinguishing features.

Mathematical Process:
1. STFT: Apply windowed Fourier transform to extract frequency content over time
2. Power Spectrum: Compute magnitude squared of STFT output
3. Mel Filterbank: Apply triangular mel-scale filters to compress frequency representation
4. Log Compression: Convert to decibel scale to normalize dynamic range
5. Normalization: Scale values to 0-1 range for neural network input

Spectrogram Parameters:
• n_mels: 128 frequency bins
• n_fft: 2048 samples (approximately 93ms window at 22,050 Hz)
• hop_length: 512 samples (approximately 23ms between frames)
• fmin: 150 Hz (below most bird calls)
• fmax: 15,000 Hz (above most bird calls)
• Output shape: (128, 216) per 5-second chunk
""",
    
    "model_architecture": """
The classification model is based on EfficientNet-B0, a state-of-the-art convolutional neural network architecture developed by Google Research.

Rationale for EfficientNet:
• Achieves excellent accuracy with fewer parameters than alternatives
• Pretrained on ImageNet enables transfer learning
• Compound scaling balances network depth, width, and resolution
• Mobile Inverted Bottleneck Convolution (MBConv) blocks provide efficiency
• Squeeze-and-Excitation (SE) attention mechanism highlights important features

Architecture Details:
• Input: (3, 128, 216) - 3-channel spectrogram (replicated for RGB compatibility)
• Backbone: EfficientNet-B0 pretrained on ImageNet
• Feature dimension: 1,280 after global average pooling
• Classifier: Dropout (0.3) → Dense (54 classes) → Softmax
• Total parameters: Approximately 4 million
• All parameters trainable (full fine-tuning)
""",
    
    "training": """
The model was trained using the following configuration:

Optimizer: AdamW
• Learning rate: 0.001 with cosine annealing warm restarts
• Weight decay: 0.0001 (L2 regularization)
• Beta parameters: β₁ = 0.9, β₂ = 0.999

Loss Function: Cross-Entropy with Label Smoothing
• Label smoothing factor: 0.1
• Prevents overconfident predictions

Learning Rate Schedule: Cosine Annealing with Warm Restarts
• T₀ = 10 epochs
• T_mult = 2
• Minimum learning rate: 1e-6

Regularization Techniques:
• Dropout: 0.3 in classifier layer
• Label smoothing: 0.1
• Weight decay: 0.0001
• Early stopping: Patience of 7 epochs

Data Augmentation (Applied Online During Training):
• Time masking: Random 0-30 frames set to zero
• Frequency masking: Random 0-20 mel bands set to zero
• Gaussian noise: Added with σ = 0.01

Training Environment:
• GPU: NVIDIA GeForce RTX 3050 (4GB VRAM)
• Mixed precision: FP16 (Automatic Mixed Precision)
• Batch size: 16
• Total training time: Approximately 2 hours
• Epochs: 30 (early stopped at epoch 17)
• Best model saved at epoch 10
""",
    
    "explainability": """
Gradient-weighted Class Activation Mapping (Grad-CAM) was implemented to provide visual explanations for model predictions.

Algorithm:
1. Forward pass: Compute feature maps from the last convolutional layer and obtain prediction
2. Backward pass: Compute gradients of the target class score with respect to feature maps
3. Weight calculation: Global average pool the gradients to obtain importance weights
4. Heatmap generation: Compute weighted combination of feature maps
5. ReLU activation: Keep only positive contributions
6. Upsampling: Resize heatmap to match input spectrogram dimensions

Mathematical Formulation:
• Importance weights: α_k = (1/Z) × Σᵢ Σⱼ (∂y^c / ∂A^k_ij)
• Grad-CAM heatmap: L^c = ReLU(Σ_k α_k × A^k)

Where:
• y^c is the score for target class c
• A^k is the activation map of channel k
• Z is the number of pixels in the activation map

Interpretation:
• Red/Yellow regions indicate high model attention (likely bird call)
• Blue regions indicate low attention (background noise)
• Researchers can verify if the model focuses on relevant acoustic features
"""
}

RESULTS = {
    "training_performance": """
The model was trained for 17 epochs before early stopping was triggered. Training progress showed rapid initial learning followed by gradual improvement.

Training History Summary:
• Epoch 1: Train Accuracy 45.23%, Validation Accuracy 38.52%
• Epoch 5: Train Accuracy 96.86%, Validation Accuracy 64.42%
• Epoch 10: Train Accuracy 99.83%, Validation Accuracy 68.79% (Best Model)
• Epoch 17: Train Accuracy 98.25%, Validation Accuracy 62.79% (Early Stopping)

The gap between training accuracy (99.83%) and validation accuracy (68.79%) indicates some degree of overfitting, which is discussed in the limitations section.
""",
    
    "test_evaluation": """
The final model was evaluated on the held-out test set of 5,404 chunks from 287 recordings.

Overall Performance Metrics:
• Top-1 Accuracy: 72.26%
• Top-3 Accuracy: 82.83%
• Top-5 Accuracy: 85.55%
• F1 Score (Macro): 63.21%
• F1 Score (Weighted): 72.26%

The Top-5 accuracy of 85.55% indicates that the correct species is among the top 5 predictions in the vast majority of cases, which is practically useful for researchers who can quickly verify from a short list.
""",
    
    "per_species": """
Performance varied across species, with some achieving over 90% F1 score while others performed less well.

Top 5 Performing Species:
1. Selasphorus platycercus (Broad-tailed Hummingbird): F1 = 0.92
2. Botaurus lentiginosus (American Bittern): F1 = 0.91
3. Coccyzus erythropthalmus (Black-billed Cuckoo): F1 = 0.91
4. Cyanocitta cristata (Blue Jay): F1 = 0.91
5. Anthus rubescens (American Pipit): F1 = 0.88

Species with distinctive, consistent vocalizations tended to perform better, while species with variable calls or limited training samples showed lower performance.
""",
    
    "real_world": """
To validate real-world performance, the system was tested on recordings not present in the training data.

Real-World Test Results:
• Recording XC475302: Actual = Blue-winged Warbler, Predicted = Blue-winged Warbler (97.0% confidence) ✓
• Recording XC416747: Actual = Blue-winged Warbler, Predicted = Blue-winged Warbler (96.6% confidence) ✓
• Recording XC481834: Actual = Blue-winged Warbler, Predicted = Bald Eagle (95.3% confidence) ✗

Real-world accuracy: 66.7% (2 out of 3 correct)

The misclassification of XC481834 with high confidence (95.3%) highlights an important limitation: the model can be "confidently wrong." This motivated the implementation of uncertainty warnings and the 80% confidence threshold for positive detection.
"""
}

DISCUSSION = {
    "strengths": """
The developed system demonstrates several strengths:

1. Reasonable Accuracy: Achieving 72.26% top-1 and 85.55% top-5 accuracy on a 54-class problem with limited training data is competitive with similar systems in the literature.

2. Explainability: The Grad-CAM visualizations provide valuable insight into model decision-making, allowing researchers to verify that predictions are based on relevant acoustic features rather than artifacts or noise.

3. Confidence Awareness: The 80% threshold for positive detection and the reliability scoring system help users understand when predictions can be trusted.

4. Practical Interface: The Streamlit web application makes the system accessible to researchers without programming expertise.

5. Real-world Validation: Testing on completely unseen recordings demonstrated that the model can generalize beyond the training distribution, correctly identifying 2 out of 3 test cases.

6. Professional Output: The PDF report generation feature provides documentation suitable for research publications and field reports.
""",
    
    "limitations": """
Several limitations should be acknowledged:

1. Overfitting: The 31% gap between training accuracy (99.83%) and validation accuracy (68.79%) indicates the model has memorized some training-specific patterns. This likely results from:
   • Limited training data (only 1,926 recordings)
   • High model capacity (4 million parameters)
   • Insufficient data augmentation

2. Confidence Calibration: The model can predict with very high confidence (>95%) on incorrect species, as demonstrated by the XC481834 misclassification. This occurs because:
   • Softmax outputs tend to be overconfident
   • The model has not been exposed to sufficient diversity in recording conditions

3. Single-label Classification: The current system predicts only one species per 5-second chunk, but real forest recordings often contain multiple species calling simultaneously.

4. Incomplete Dataset: Only 46% of downloaded recordings (1,926 out of 4,521) were usable due to metadata scraping failures, potentially limiting the diversity of training examples.

5. Class Imbalance: Training data ranges from 98 to 1,253 samples per species (12.8x ratio), which may bias predictions toward more common species.

6. Recording Variability: The model may perform differently on recordings made with different equipment, at different distances, or in different environmental conditions than the training data.
"""
}

FUTURE_WORK = """
Several directions for future improvement have been identified:

Short-term Improvements:
• Complete metadata scraping for remaining 2,205 recordings to expand training data
• Implement stronger data augmentation (SpecAugment, Mixup, CutMix)
• Apply confidence calibration using temperature scaling or Platt scaling
• Add more uncertainty quantification metrics

Medium-term Improvements:
• Convert to multi-label classification to detect multiple species per chunk
• Implement transfer learning from BirdNET embeddings for improved features
• Build ensemble models combining multiple architectures for robustness
• Add species range validation using geographic metadata

Long-term Vision:
• Develop real-time detection capability for streaming audio input
• Create mobile applications for iOS and Android field deployment
• Implement active learning with user feedback for continuous improvement
• Integrate with camera systems for multi-modal bird identification
"""

CONCLUSION = """
This project successfully developed a confidence-aware, explainable bird species detection system using bioacoustic signals and deep learning.

Key Contributions:

1. Working End-to-End System: A complete pipeline from audio input to species prediction with visual explanation, accessible through an intuitive web interface.

2. Explainability: Grad-CAM visualizations show which spectrogram regions influence predictions, enabling researchers to verify and understand model decisions.

3. Confidence Awareness: The 80% threshold for positive detection and reliability scoring system help users understand prediction trustworthiness, addressing the problem of overconfident incorrect predictions.

4. Practical Validation: Real-world testing on unseen recordings demonstrated both the capabilities and limitations of the system, providing honest assessment of performance.

Key Performance Metrics:
• Top-1 Accuracy: 72.26%
• Top-5 Accuracy: 85.55%
• Species Covered: 54
• Real-world Tests: 2/3 Correct (66.7%)

The system serves as a valuable tool for researchers needing quick, explainable bird species identification from audio recordings. While acknowledging its limitations through uncertainty warnings, it provides a practical solution for bioacoustic monitoring that can be improved with additional data and continued development.

The honest acknowledgment of limitations, particularly the potential for high-confidence misclassifications, represents an important contribution toward responsible deployment of AI systems in scientific applications where trust and verification are essential.
"""

REFERENCES = [
    "Tan, M., & Le, Q. (2019). EfficientNet: Rethinking Model Scaling for Convolutional Neural Networks. International Conference on Machine Learning (ICML).",
    "Selvaraju, R. R., Cogswell, M., Das, A., Vedantam, R., Parikh, D., & Batra, D. (2017). Grad-CAM: Visual Explanations from Deep Networks via Gradient-based Localization. IEEE International Conference on Computer Vision (ICCV).",
    "Kahl, S., Wood, C. M., Eibl, M., & Klinck, H. (2021). BirdNET: A deep learning solution for avian diversity monitoring. Ecological Informatics, 61, 101236.",
    "Kong, Q., Cao, Y., Iqbal, T., Wang, Y., Wang, W., & Plumbley, M. D. (2020). PANNs: Large-Scale Pretrained Audio Neural Networks for Audio Pattern Recognition. IEEE/ACM Transactions on Audio, Speech, and Language Processing, 28, 2880-2894.",
    "McFee, B., Raffel, C., Liang, D., Ellis, D. P., McVicar, M., Battenberg, E., & Nieto, O. (2015). librosa: Audio and Music Signal Analysis in Python. Proceedings of the 14th Python in Science Conference, 18-25.",
    "Xeno-Canto Foundation. (2024). Xeno-Canto: Sharing Bird Sounds from Around the World. https://xeno-canto.org",
    "Park, D. S., Chan, W., Zhang, Y., Chiu, C. C., Zoph, B., Cubuk, E. D., & Le, Q. V. (2019). SpecAugment: A Simple Data Augmentation Method for Automatic Speech Recognition. Interspeech 2019.",
    "Loshchilov, I., & Hutter, F. (2019). Decoupled Weight Decay Regularization. International Conference on Learning Representations (ICLR).",
]

SPECIES_LIST = [
    ("Myiarchus cinerascens", "Ash-throated Flycatcher", 1253),
    ("Spizella breweri", "Brewer's Sparrow", 1249),
    ("Spinus tristis", "American Goldfinch", 1062),
    ("Artemisiospiza belli", "Bell's Sparrow", 895),
    ("Polioptila caerulea", "Blue-gray Gnatcatcher", 827),
    ("Setophaga striata", "Blackpoll Warbler", 778),
    ("Calypte anna", "Anna's Hummingbird", 720),
    ("Empidonax alnorum", "Alder Flycatcher", 711),
    ("Dolichonyx oryzivorus", "Bobolink", 706),
    ("Toxostoma rufum", "Brown Thrasher", 695),
    ("Spizelloides arborea", "American Tree Sparrow", 585),
    ("Thryomanes bewickii", "Bewick's Wren", 574),
    ("Mniotilta varia", "Black-and-white Warbler", 552),
    ("Corvus brachyrhynchos", "American Crow", 549),
    ("Setophaga virens", "Black-throated Green Warbler", 517),
    ("Scolopax minor", "American Woodcock", 513),
    ("Setophaga ruticilla", "American Redstart", 499),
    ("Riparia riparia", "Bank Swallow", 497),
    ("Turdus migratorius", "American Robin", 482),
    ("Selasphorus platycercus", "Broad-tailed Hummingbird", 454),
    ("Cyanocitta cristata", "Blue Jay", 452),
    ("Poecile atricapillus", "Black-capped Chickadee", 446),
    ("Setophaga nigrescens", "Black-throated Gray Warbler", 435),
    ("Setophaga fusca", "Blackburnian Warbler", 434),
    ("Certhia americana", "Brown Creeper", 434),
    ("Molothrus ater", "Brown-headed Cowbird", 433),
    ("Sayornis nigricans", "Black Phoebe", 414),
    ("Icterus galbula", "Baltimore Oriole", 412),
    ("Amphispiza bilineata", "Black-throated Sparrow", 400),
    ("Pheucticus melanocephalus", "Black-headed Grosbeak", 391),
    ("Buteo platypterus", "Broad-winged Hawk", 367),
    ("Passerina caerulea", "Blue Grosbeak", 329),
    ("Vermivora cyanoptera", "Blue-winged Warbler", 326),
    ("Megaceryle alcyon", "Belted Kingfisher", 325),
    ("Vireo solitarius", "Blue-headed Vireo", 320),
    ("Strix varia", "Barred Owl", 318),
    ("Icterus bullockii", "Bullock's Oriole", 296),
    ("Archilochus alexandri", "Black-chinned Hummingbird", 278),
    ("Setophaga caerulescens", "Black-throated Blue Warbler", 274),
    ("Psaltriparus minimus", "Bushtit", 271),
    ("Haliaeetus leucocephalus", "Bald Eagle", 250),
    ("Pica hudsonia", "Black-billed Magpie", 235),
    ("Spatula discors", "Blue-winged Teal", 224),
    ("Coccyzus erythropthalmus", "Black-billed Cuckoo", 201),
    ("Hirundo rustica", "Barn Swallow", 200),
    ("Botaurus lentiginosus", "American Bittern", 199),
    ("Chroicocephalus philadelphia", "Bonaparte's Gull", 192),
    ("Mareca americana", "American Wigeon", 165),
    ("Anthus rubescens", "American Pipit", 140),
    ("Euphagus cyanocephalus", "Brewer's Blackbird", 140),
    ("Calidris bairdii", "Baird's Sandpiper", 109),
    ("Falco sparverius", "American Kestrel", 107),
    ("Recurvirostra americana", "American Avocet", 106),
    ("Bucephala albeola", "Bufflehead", 98),
]


# ============================================================
# DOCX GENERATION FUNCTIONS
# ============================================================

def create_docx_report():
    """Create the complete DOCX report"""
    
    print("📝 Creating DOCX report...")
    
    doc = Document()
    
    # ========== DOCUMENT STYLES ==========
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Title style
    title_style = doc.styles.add_style('ReportTitle', WD_STYLE_TYPE.PARAGRAPH)
    title_style.font.name = 'Times New Roman'
    title_style.font.size = Pt(16)
    title_style.font.bold = True
    title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_style.paragraph_format.space_after = Pt(6)
    
    # Heading 1 style
    h1_style = doc.styles['Heading 1']
    h1_style.font.name = 'Times New Roman'
    h1_style.font.size = Pt(14)
    h1_style.font.bold = True
    h1_style.font.color.rgb = RGBColor(0, 0, 0)
    h1_style.paragraph_format.space_before = Pt(18)
    h1_style.paragraph_format.space_after = Pt(6)
    
    # Heading 2 style
    h2_style = doc.styles['Heading 2']
    h2_style.font.name = 'Times New Roman'
    h2_style.font.size = Pt(12)
    h2_style.font.bold = True
    h2_style.font.color.rgb = RGBColor(0, 0, 0)
    h2_style.paragraph_format.space_before = Pt(12)
    h2_style.paragraph_format.space_after = Pt(6)
    
    # Caption style
    caption_style = doc.styles.add_style('FigureCaption', WD_STYLE_TYPE.PARAGRAPH)
    caption_style.font.name = 'Times New Roman'
    caption_style.font.size = Pt(10)
    caption_style.font.italic = True
    caption_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_style.paragraph_format.space_before = Pt(6)
    caption_style.paragraph_format.space_after = Pt(12)
    
    # ========== TITLE PAGE ==========
    
    # Add some space at top
    for _ in range(6):
        doc.add_paragraph()
    
    # Title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(REPORT_TITLE)
    title_run.bold = True
    title_run.font.size = Pt(18)
    title_run.font.name = 'Times New Roman'
    
    doc.add_paragraph()
    
    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run("Technical Report")
    sub_run.font.size = Pt(14)
    sub_run.font.name = 'Times New Roman'
    
    # Add space
    for _ in range(4):
        doc.add_paragraph()
    
    # Author info
    author_para = doc.add_paragraph()
    author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author_para.add_run(f"Author: {AUTHOR_NAME}\n").font.size = Pt(12)
    author_para.add_run(f"Institution: {INSTITUTION}\n").font.size = Pt(12)
    author_para.add_run(f"Date: {DATE}").font.size = Pt(12)
    
    # Page break
    doc.add_page_break()
    
    # ========== ABSTRACT ==========
    
    doc.add_heading("Abstract", level=1)
    
    for para in ABSTRACT.strip().split('\n\n'):
        p = doc.add_paragraph(para.strip())
        p.paragraph_format.first_line_indent = Inches(0.5)
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Keywords
    doc.add_paragraph()
    keywords_para = doc.add_paragraph()
    keywords_run = keywords_para.add_run("Keywords: ")
    keywords_run.bold = True
    keywords_para.add_run(KEYWORDS)
    
    doc.add_page_break()
    
    # ========== TABLE OF CONTENTS PLACEHOLDER ==========
    
    doc.add_heading("Table of Contents", level=1)
    doc.add_paragraph("[Table of Contents - Update after finalizing document]")
    doc.add_paragraph()
    doc.add_paragraph("1. Introduction")
    doc.add_paragraph("2. Methodology")
    doc.add_paragraph("3. Results")
    doc.add_paragraph("4. Discussion")
    doc.add_paragraph("5. Future Work")
    doc.add_paragraph("6. Conclusion")
    doc.add_paragraph("7. References")
    doc.add_paragraph("8. Appendix")
    
    doc.add_page_break()
    
    # ========== 1. INTRODUCTION ==========
    
    doc.add_heading("1. Introduction", level=1)
    
    doc.add_heading("1.1 Background", level=2)
    add_justified_paragraph(doc, INTRODUCTION["background"])
    
    doc.add_heading("1.2 Problem Statement", level=2)
    add_justified_paragraph(doc, INTRODUCTION["problem_statement"])
    
    doc.add_heading("1.3 Objectives", level=2)
    add_justified_paragraph(doc, INTRODUCTION["objectives"])
    
    doc.add_heading("1.4 Scope", level=2)
    add_justified_paragraph(doc, INTRODUCTION["scope"])
    
    doc.add_page_break()
    
    # ========== 2. METHODOLOGY ==========
    
    doc.add_heading("2. Methodology", level=1)
    
    doc.add_heading("2.1 System Architecture", level=2)
    add_justified_paragraph(doc, METHODOLOGY["system_architecture"])
    
    # Figure 1 placeholder
    add_figure_placeholder(doc, 1, "System Architecture",
        "Pipeline diagram showing: Audio Input → Mel-Spectrogram → EfficientNet Classifier → Species Predictions + Grad-CAM Heatmap")
    
    doc.add_heading("2.2 Dataset", level=2)
    add_justified_paragraph(doc, METHODOLOGY["dataset"])
    
    # Dataset statistics table
    doc.add_paragraph()
    table_data = [
        ["Metric", "Value"],
        ["Total recordings downloaded", "4,521"],
        ["Usable recordings", "1,926"],
        ["Species covered", "54"],
        ["Total audio chunks", "34,811"],
        ["Training chunks", "23,839 (70%)"],
        ["Validation chunks", "5,568 (15%)"],
        ["Test chunks", "5,404 (15%)"],
    ]
    add_table(doc, table_data, "Table 1: Dataset Statistics")
    
    doc.add_heading("2.3 Audio Preprocessing", level=2)
    add_justified_paragraph(doc, METHODOLOGY["preprocessing"])
    
    doc.add_heading("2.4 Feature Extraction (Mel-Spectrogram)", level=2)
    add_justified_paragraph(doc, METHODOLOGY["feature_extraction"])
    
    # Spectrogram parameters table
    table_data = [
        ["Parameter", "Value", "Purpose"],
        ["n_mels", "128", "Frequency bins"],
        ["n_fft", "2048", "FFT window size (~93ms)"],
        ["hop_length", "512", "Frame hop (~23ms)"],
        ["fmin", "150 Hz", "Minimum frequency"],
        ["fmax", "15,000 Hz", "Maximum frequency"],
        ["Output shape", "(128, 216)", "Per 5-second chunk"],
    ]
    add_table(doc, table_data, "Table 2: Mel-Spectrogram Parameters")
    
    # Figure 2 placeholder
    add_figure_placeholder(doc, 2, "Mel-Spectrogram Conversion",
        "Side-by-side comparison: LEFT - Audio waveform (amplitude vs time), RIGHT - Mel-spectrogram (frequency vs time) with bird call visible as bright pattern")
    
    doc.add_heading("2.5 Model Architecture", level=2)
    add_justified_paragraph(doc, METHODOLOGY["model_architecture"])
    
    # Model architecture table
    table_data = [
        ["Component", "Details"],
        ["Backbone", "EfficientNet-B0 (pretrained on ImageNet)"],
        ["Feature dimension", "1,280"],
        ["Classifier", "Dropout(0.3) → Dense(54) → Softmax"],
        ["Total parameters", "~4 million"],
        ["Input shape", "(3, 128, 216)"],
        ["Output", "54 class probabilities"],
    ]
    add_table(doc, table_data, "Table 3: Model Architecture Summary")
    
    doc.add_heading("2.6 Training Configuration", level=2)
    add_justified_paragraph(doc, METHODOLOGY["training"])
    
    # Training config table
    table_data = [
        ["Parameter", "Value"],
        ["Optimizer", "AdamW"],
        ["Learning Rate", "0.001 (cosine annealing)"],
        ["Weight Decay", "0.0001"],
        ["Batch Size", "16"],
        ["Epochs", "30 (early stopped at 17)"],
        ["Loss Function", "CrossEntropy + Label Smoothing (0.1)"],
        ["Mixed Precision", "FP16"],
        ["GPU", "NVIDIA RTX 3050 (4GB)"],
        ["Training Time", "~2 hours"],
    ]
    add_table(doc, table_data, "Table 4: Training Configuration")
    
    doc.add_heading("2.7 Explainability: Grad-CAM", level=2)
    add_justified_paragraph(doc, METHODOLOGY["explainability"])
    
    doc.add_page_break()
    
    # ========== 3. RESULTS ==========
    
    doc.add_heading("3. Results", level=1)
    
    doc.add_heading("3.1 Training Performance", level=2)
    add_justified_paragraph(doc, RESULTS["training_performance"])
    
    # Training history table
    table_data = [
        ["Epoch", "Train Loss", "Train Acc", "Val Loss", "Val Acc", "Notes"],
        ["1", "2.4521", "45.23%", "2.8934", "38.52%", "Initial"],
        ["5", "1.1442", "96.86%", "2.5174", "64.42%", "Rapid learning"],
        ["10", "1.0051", "99.83%", "2.3091", "68.79%", "Best Model ✓"],
        ["17", "1.0646", "98.25%", "2.6662", "62.79%", "Early Stopping"],
    ]
    add_table(doc, table_data, "Table 5: Training History")
    
    # Figure 3 placeholder
    add_figure_placeholder(doc, 3, "Training Curves",
        "Line graph showing Train Accuracy (rising to ~99%), Val Accuracy (plateaus at ~68%), Train Loss (decreasing), Val Loss (U-shaped). Mark best model at Epoch 10. Source: TensorBoard logs in 05_Model/Training_Logs/")
    
    doc.add_heading("3.2 Test Set Evaluation", level=2)
    add_justified_paragraph(doc, RESULTS["test_evaluation"])
    
    # Metrics table
    table_data = [
        ["Metric", "Value"],
        ["Top-1 Accuracy", "72.26%"],
        ["Top-3 Accuracy", "82.83%"],
        ["Top-5 Accuracy", "85.55%"],
        ["F1 Score (Macro)", "63.21%"],
        ["F1 Score (Weighted)", "72.26%"],
        ["Test Samples", "5,404"],
    ]
    add_table(doc, table_data, "Table 6: Test Set Performance Metrics")
    
    # Figure 4 placeholder
    add_figure_placeholder(doc, 4, "Confusion Matrix",
        "54×54 normalized confusion matrix heatmap. X-axis: Predicted Species, Y-axis: True Species. Strong diagonal indicates good performance. Source: 07_Evaluation/Confusion_Matrices/confusion_matrix.png")
    
    doc.add_heading("3.3 Per-Species Performance", level=2)
    add_justified_paragraph(doc, RESULTS["per_species"])
    
    # Top species table
    table_data = [
        ["Species", "English Name", "Precision", "Recall", "F1"],
        ["Selasphorus platycercus", "Broad-tailed Hummingbird", "0.95", "0.89", "0.92"],
        ["Botaurus lentiginosus", "American Bittern", "0.89", "0.94", "0.91"],
        ["Coccyzus erythropthalmus", "Black-billed Cuckoo", "0.95", "0.88", "0.91"],
        ["Cyanocitta cristata", "Blue Jay", "0.87", "0.94", "0.91"],
        ["Anthus rubescens", "American Pipit", "0.88", "0.89", "0.88"],
    ]
    add_table(doc, table_data, "Table 7: Top 5 Performing Species")
    
    doc.add_heading("3.4 Explainability Results", level=2)
    
    p = doc.add_paragraph()
    p.add_run("Grad-CAM visualizations demonstrate where the model focuses when making predictions. ").font.size = Pt(12)
    p.add_run("These visualizations enable researchers to verify that the model is attending to relevant acoustic features.").font.size = Pt(12)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Figure 5 placeholder
    add_figure_placeholder(doc, 5, "Grad-CAM Example: Correct Prediction",
        "Three-panel visualization showing: (1) Original Spectrogram, (2) Grad-CAM Heatmap, (3) Overlay. Caption: Blue-winged Warbler correctly identified with 97% confidence. Model focuses on bird call frequency bands. Source: 06_Explainability/GradCAM/sample_visualizations/")
    
    # Figure 6 placeholder
    add_figure_placeholder(doc, 6, "Grad-CAM Example: Incorrect Prediction",
        "Three-panel visualization for recording XC481834. Actual: Blue-winged Warbler. Predicted: Bald Eagle (95.3%). Shows model focusing on unexpected features, explaining the misclassification.")
    
    doc.add_heading("3.5 Real-World Validation", level=2)
    add_justified_paragraph(doc, RESULTS["real_world"])
    
    # Real-world test table
    table_data = [
        ["Recording", "Actual Species", "Prediction", "Confidence", "Result"],
        ["XC475302", "Blue-winged Warbler", "Blue-winged Warbler", "97.0%", "✓ Correct"],
        ["XC416747", "Blue-winged Warbler", "Blue-winged Warbler", "96.6%", "✓ Correct"],
        ["XC481834", "Blue-winged Warbler", "Bald Eagle", "95.3%", "✗ Wrong"],
    ]
    add_table(doc, table_data, "Table 8: Real-World Validation Results")
    
    doc.add_heading("3.6 Deployed Application", level=2)
    
    p = doc.add_paragraph()
    p.add_run("A web interface was developed using Streamlit to make the system accessible to researchers without programming expertise. ").font.size = Pt(12)
    p.add_run("The application provides audio upload, species search, visualization, and report generation capabilities.").font.size = Pt(12)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Features table
    table_data = [
        ["Feature", "Description"],
        ["Audio Upload", "Supports WAV, MP3, OGG, FLAC formats"],
        ["Target Species Search", "Search for specific species"],
        ["All Species Detection", "Detect all 54 species"],
        ["Confidence Display", "Color-coded (🟢 High, 🟡 Medium, 🔴 Low)"],
        ["80% Threshold", "Species marked 'Present' only if ≥80%"],
        ["Reliability Warnings", "Alerts for unusual prediction patterns"],
        ["Detection Timeline", "Shows when species detected in recording"],
        ["Grad-CAM Visualization", "Visual explanation of predictions"],
        ["PDF Export", "Download professional report"],
    ]
    add_table(doc, table_data, "Table 9: Application Features")
    
    # Figure 7 placeholder
    add_figure_placeholder(doc, 7, "Streamlit Application Interface",
        "Screenshot of the web application showing: Header, Audio uploader, Target species dropdown, Detection result panel (✅ SPECIES PRESENT or ❌ NOT DETECTED), Confidence score, Reliability warnings, Visualization tabs. Take screenshot from running application.")
    
    # Figure 8 placeholder
    add_figure_placeholder(doc, 8, "PDF Report Output",
        "First page of generated PDF report showing: Title, Audio information, Target species result, All detected species table. Download a PDF from the application to use as image.")
    
    doc.add_page_break()
    
    # ========== 4. DISCUSSION ==========
    
    doc.add_heading("4. Discussion", level=1)
    
    doc.add_heading("4.1 Strengths", level=2)
    add_justified_paragraph(doc, DISCUSSION["strengths"])
    
    doc.add_heading("4.2 Limitations", level=2)
    add_justified_paragraph(doc, DISCUSSION["limitations"])
    
    doc.add_page_break()
    
    # ========== 5. FUTURE WORK ==========
    
    doc.add_heading("5. Future Work", level=1)
    add_justified_paragraph(doc, FUTURE_WORK)
    
    doc.add_page_break()
    
    # ========== 6. CONCLUSION ==========
    
    doc.add_heading("6. Conclusion", level=1)
    add_justified_paragraph(doc, CONCLUSION)
    
    # Summary metrics table
    table_data = [
        ["Metric", "Value"],
        ["Top-1 Accuracy", "72.26%"],
        ["Top-5 Accuracy", "85.55%"],
        ["Species Covered", "54"],
        ["Real-world Tests", "2/3 Correct (66.7%)"],
    ]
    add_table(doc, table_data, "Table 10: Summary of Key Performance Metrics")
    
    doc.add_page_break()
    
    # ========== 7. REFERENCES ==========
    
    doc.add_heading("7. References", level=1)
    
    for i, ref in enumerate(REFERENCES, 1):
        p = doc.add_paragraph()
        p.add_run(f"[{i}] ").bold = True
        p.add_run(ref)
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)
        p.paragraph_format.space_after = Pt(6)
    
    doc.add_page_break()
    
    # ========== 8. APPENDIX ==========
    
    doc.add_heading("8. Appendix", level=1)
    
    doc.add_heading("A. Complete Species List (54 Species)", level=2)
    
    # Create species table
    species_table_data = [["#", "Scientific Name", "English Name", "Training Samples"]]
    for i, (sci, eng, samples) in enumerate(SPECIES_LIST, 1):
        species_table_data.append([str(i), sci, eng, str(samples)])
    
    # Add table in chunks to avoid page breaks
    add_table(doc, species_table_data[:28], "Table A1: Species List (1-27)")
    doc.add_paragraph()
    add_table(doc, [species_table_data[0]] + species_table_data[28:], "Table A2: Species List (28-54)")
    
    doc.add_heading("B. Software and Hardware", level=2)
    
    # Software table
    table_data = [
        ["Component", "Technology/Version"],
        ["Programming Language", "Python 3.11"],
        ["Deep Learning", "PyTorch 2.10, timm"],
        ["Audio Processing", "librosa 0.10"],
        ["Web Interface", "Streamlit 1.x"],
        ["Visualization", "Matplotlib, OpenCV"],
        ["PDF Generation", "ReportLab"],
    ]
    add_table(doc, table_data, "Table B1: Software Stack")
    
    # Hardware table
    table_data = [
        ["Component", "Specification"],
        ["GPU", "NVIDIA GeForce RTX 3050 (4GB VRAM)"],
        ["CUDA Version", "12.7"],
        ["Training Time", "~2 hours"],
        ["Inference Time", "~0.5 seconds per 5-second chunk"],
    ]
    add_table(doc, table_data, "Table B2: Hardware Configuration")
    
    # ========== SAVE DOCUMENT ==========
    
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    docx_path = OUTPUT_DIR / "Bird_Detection_Technical_Report.docx"
    doc.save(str(docx_path))
    
    print(f"✅ DOCX saved: {docx_path}")
    return docx_path


def add_justified_paragraph(doc, text):
    """Add justified paragraph(s) to document"""
    paragraphs = text.strip().split('\n\n')
    for para_text in paragraphs:
        if para_text.strip():
            # Handle bullet points
            if para_text.strip().startswith('•'):
                lines = para_text.strip().split('\n')
                for line in lines:
                    if line.strip():
                        p = doc.add_paragraph(line.strip(), style='List Bullet')
            else:
                p = doc.add_paragraph(para_text.strip())
                p.paragraph_format.first_line_indent = Inches(0.5)
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.space_after = Pt(6)


def add_table(doc, data, caption=None):
    """Add a formatted table to document"""
    table = doc.add_table(rows=len(data), cols=len(data[0]))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for i, row_data in enumerate(data):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            cell.text = str(cell_text)
            
            # Header row formatting
            if i == 0:
                cell.paragraphs[0].runs[0].bold = True
                # Set background color for header
                shading = OxmlElement('w:shd')
                shading.set(qn('w:fill'), 'E6E6E6')
                cell._tc.get_or_add_tcPr().append(shading)
            
            # Center align
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add caption
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run(caption)
        run.italic = True
        run.font.size = Pt(10)


def add_figure_placeholder(doc, num, title, description):
    """Add a figure placeholder box to document"""
    
    # Create a bordered placeholder
    doc.add_paragraph()
    
    # Top border line
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("─" * 60)
    
    # Placeholder content
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"\n\n\nFIGURE {num}\n\n").bold = True
    p.add_run(f"{title}\n\n").font.size = Pt(11)
    p.add_run(f"[{description}]\n\n\n").font.size = Pt(10)
    
    # Bottom border line
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("─" * 60)
    
    # Caption
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(f"Figure {num}: {title}")
    run.italic = True
    run.font.size = Pt(10)
    
    doc.add_paragraph()


# ============================================================
# PDF GENERATION FUNCTIONS
# ============================================================

def create_pdf_report():
    """Create the complete PDF report"""
    
    print("📄 Creating PDF report...")
    
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    pdf_path = OUTPUT_DIR / "Bird_Detection_Technical_Report.pdf"
    
    doc = SimpleDocTemplate(
        str(pdf_path),
        pagesize=A4,
        rightMargin=72,
        leftMargin=72,
        topMargin=72,
        bottomMargin=72
    )
    
    # Styles
    styles = getSampleStyleSheet()
    
    # Custom styles - use unique names to avoid conflicts
    try:
        styles.add(ParagraphStyle(
            name='ReportTitle',
            parent=styles['Heading1'],
            fontSize=18,
            spaceAfter=20,
            alignment=TA_CENTER,
            fontName='Times-Bold'
        ))
    except:
        pass
    
    try:
        styles.add(ParagraphStyle(
            name='SectionHeading',
            parent=styles['Heading1'],
            fontSize=14,
            spaceBefore=20,
            spaceAfter=10,
            fontName='Times-Bold'
        ))
    except:
        pass
    
    try:
        styles.add(ParagraphStyle(
            name='SubsectionHeading',
            parent=styles['Heading2'],
            fontSize=12,
            spaceBefore=15,
            spaceAfter=8,
            fontName='Times-Bold'
        ))
    except:
        pass
    
    try:
        styles.add(ParagraphStyle(
            name='BodyText',
            parent=styles['Normal'],
            fontSize=11,
            leading=14,
            alignment=TA_JUSTIFY,
            fontName='Times-Roman',
            firstLineIndent=36,
            spaceAfter=6
        ))
    except:
        pass
    
    try:
        styles.add(ParagraphStyle(
            name='FigureCaption',
            parent=styles['Normal'],
            fontSize=10,
            alignment=TA_CENTER,
            fontName='Times-Italic',
            spaceBefore=6,
            spaceAfter=12
        ))
    except:
        pass
    
    try:
        styles.add(ParagraphStyle(
            name='TableCaption',
            parent=styles['Normal'],
            fontSize=10,
            alignment=TA_CENTER,
            fontName='Times-Italic',
            spaceBefore=6,
            spaceAfter=12
        ))
    except:
        pass
    
    try:
        styles.add(ParagraphStyle(
            name='CenteredText',
            parent=styles['Normal'],
            fontSize=12,
            alignment=TA_CENTER,
            fontName='Times-Roman'
        ))
    except:
        pass
    
    # Image folder path
    image_folder = BASE_DIR / "10_Outputs" / "Report_images"
    image_folder.mkdir(parents=True, exist_ok=True)
    
    # Build content
    story = []
    
    # ========== TITLE PAGE ==========
    
    story.append(Spacer(1, 2*inch))
    story.append(Paragraph(REPORT_TITLE, styles['ReportTitle']))
    story.append(Spacer(1, 0.5*inch))
    story.append(Paragraph("Technical Report", styles['CenteredText']))
    story.append(Spacer(1, 1*inch))
    story.append(Paragraph(f"Author: {AUTHOR_NAME}", styles['CenteredText']))
    story.append(Paragraph(f"Institution: {INSTITUTION}", styles['CenteredText']))
    story.append(Paragraph(f"Date: {DATE}", styles['CenteredText']))
    story.append(PageBreak())
    
    # ========== ABSTRACT ==========
    
    story.append(Paragraph("Abstract", styles['SectionHeading']))
    
    for para in ABSTRACT.strip().split('\n\n'):
        if para.strip():
            story.append(Paragraph(para.strip(), styles['BodyText']))
    
    story.append(Spacer(1, 12))
    story.append(Paragraph(f"<b>Keywords:</b> {KEYWORDS}", styles['Normal']))
    story.append(PageBreak())
    
    # ========== 1. INTRODUCTION ==========
    
    story.append(Paragraph("1. Introduction", styles['SectionHeading']))
    
    story.append(Paragraph("1.1 Background", styles['SubsectionHeading']))
    add_pdf_paragraphs(story, INTRODUCTION["background"], styles)
    
    story.append(Paragraph("1.2 Problem Statement", styles['SubsectionHeading']))
    add_pdf_paragraphs(story, INTRODUCTION["problem_statement"], styles)
    
    story.append(Paragraph("1.3 Objectives", styles['SubsectionHeading']))
    add_pdf_paragraphs(story, INTRODUCTION["objectives"], styles)
    
    story.append(Paragraph("1.4 Scope", styles['SubsectionHeading']))
    add_pdf_paragraphs(story, INTRODUCTION["scope"], styles)
    
    story.append(PageBreak())
    
    # ========== 2. METHODOLOGY ==========
    
    story.append(Paragraph("2. Methodology", styles['SectionHeading']))
    
    story.append(Paragraph("2.1 System Architecture", styles['SubsectionHeading']))
    add_pdf_paragraphs(story, METHODOLOGY["system_architecture"], styles)
    
    # Figure 1 placeholder
    add_pdf_figure(story, 1, "System Architecture", 
        "Pipeline diagram showing: Audio Input → Mel-Spectrogram → EfficientNet → Predictions + Grad-CAM", styles, image_folder)
    
    story.append(Paragraph("2.2 Dataset", styles['SubsectionHeading']))
    add_pdf_paragraphs(story, METHODOLOGY["dataset"], styles)
    
    # Dataset table
    table_data = [
        ["Metric", "Value"],
        ["Total recordings downloaded", "4,521"],
        ["Usable recordings", "1,926"],
        ["Species covered", "54"],
        ["Total audio chunks", "34,811"],
        ["Training chunks", "23,839 (70%)"],
        ["Validation chunks", "5,568 (15%)"],
        ["Test chunks", "5,404 (15%)"],
    ]
    add_pdf_table(story, table_data, "Table 1: Dataset Statistics", styles)
    
    story.append(Paragraph("2.3 Audio Preprocessing", styles['SubsectionHeading']))
    add_pdf_paragraphs(story, METHODOLOGY["preprocessing"], styles)
    
    story.append(Paragraph("2.4 Feature Extraction", styles['SubsectionHeading']))
    add_pdf_paragraphs(story, METHODOLOGY["feature_extraction"], styles)
    
    # Spectrogram parameters table
    table_data = [
        ["Parameter", "Value", "Purpose"],
        ["n_mels", "128", "Frequency bins"],
        ["n_fft", "2048", "FFT window (~93ms)"],
        ["hop_length", "512", "Frame hop (~23ms)"],
        ["fmin", "150 Hz", "Min frequency"],
        ["fmax", "15,000 Hz", "Max frequency"],
    ]
    add_pdf_table(story, table_data, "Table 2: Spectrogram Parameters", styles)
    
    # Figure 2 placeholder
    add_pdf_figure(story, 2, "Mel-Spectrogram Conversion",
        "Audio waveform (left) converted to mel-spectrogram (right)", styles, image_folder)
    
    story.append(Paragraph("2.5 Model Architecture", styles['SubsectionHeading']))
    add_pdf_paragraphs(story, METHODOLOGY["model_architecture"], styles)
    
    story.append(Paragraph("2.6 Training Configuration", styles['SubsectionHeading']))
    add_pdf_paragraphs(story, METHODOLOGY["training"], styles)
    
    # Training config table
    table_data = [
        ["Parameter", "Value"],
        ["Optimizer", "AdamW"],
        ["Learning Rate", "0.001"],
        ["Batch Size", "16"],
        ["Epochs", "30 (stopped at 17)"],
        ["Loss", "CrossEntropy + Label Smoothing"],
        ["GPU", "RTX 3050 (4GB)"],
    ]
    add_pdf_table(story, table_data, "Table 3: Training Configuration", styles)
    
    story.append(Paragraph("2.7 Explainability: Grad-CAM", styles['SubsectionHeading']))
    add_pdf_paragraphs(story, METHODOLOGY["explainability"], styles)
    
    story.append(PageBreak())
    
    # ========== 3. RESULTS ==========
    
    story.append(Paragraph("3. Results", styles['SectionHeading']))
    
    story.append(Paragraph("3.1 Training Performance", styles['SubsectionHeading']))
    add_pdf_paragraphs(story, RESULTS["training_performance"], styles)
    
    # Training history table
    table_data = [
        ["Epoch", "Train Acc", "Val Acc", "Notes"],
        ["1", "45.23%", "38.52%", "Initial"],
        ["5", "96.86%", "64.42%", "Rapid learning"],
        ["10", "99.83%", "68.79%", "Best Model"],
        ["17", "98.25%", "62.79%", "Early Stop"],
    ]
    add_pdf_table(story, table_data, "Table 4: Training History", styles)
    
    # Figure 3 placeholder
    add_pdf_figure(story, 3, "Training Curves",
        "Train/Val accuracy and loss over epochs", styles, image_folder)
    
    story.append(Paragraph("3.2 Test Set Evaluation", styles['SubsectionHeading']))
    add_pdf_paragraphs(story, RESULTS["test_evaluation"], styles)
    
    # Metrics table
    table_data = [
        ["Metric", "Value"],
        ["Top-1 Accuracy", "72.26%"],
        ["Top-3 Accuracy", "82.83%"],
        ["Top-5 Accuracy", "85.55%"],
        ["F1 (Weighted)", "72.26%"],
    ]
    add_pdf_table(story, table_data, "Table 5: Test Performance", styles)
    
    # Figure 4 placeholder
    add_pdf_figure(story, 4, "Confusion Matrix",
        "54×54 confusion matrix heatmap", styles, image_folder)
    
    story.append(Paragraph("3.3 Per-Species Performance", styles['SubsectionHeading']))
    add_pdf_paragraphs(story, RESULTS["per_species"], styles)
    
    story.append(Paragraph("3.4 Explainability Results", styles['SubsectionHeading']))
    story.append(Paragraph(
        "Grad-CAM visualizations demonstrate where the model focuses when making predictions, "
        "enabling researchers to verify that predictions are based on relevant acoustic features.",
        styles['BodyText']
    ))
    
    # Figure 5 placeholder
    add_pdf_figure(story, 5, "Grad-CAM: Correct Prediction",
        "Spectrogram + Heatmap + Overlay for correct classification", styles, image_folder)
    
    # Figure 6 placeholder
    add_pdf_figure(story, 6, "Grad-CAM: Incorrect Prediction",
        "Spectrogram + Heatmap + Overlay for misclassification", styles, image_folder)
    
    story.append(Paragraph("3.5 Real-World Validation", styles['SubsectionHeading']))
    add_pdf_paragraphs(story, RESULTS["real_world"], styles)
    
    # Real-world table
    table_data = [
        ["Recording", "Actual", "Predicted", "Conf.", "Result"],
        ["XC475302", "Blue-winged Warbler", "Blue-winged Warbler", "97%", "✓"],
        ["XC416747", "Blue-winged Warbler", "Blue-winged Warbler", "96.6%", "✓"],
        ["XC481834", "Blue-winged Warbler", "Bald Eagle", "95.3%", "✗"],
    ]
    add_pdf_table(story, table_data, "Table 6: Real-World Tests", styles)
    
    story.append(Paragraph("3.6 Deployed Application", styles['SubsectionHeading']))
    story.append(Paragraph(
        "A web interface was developed using Streamlit, providing audio upload, species search, "
        "visualization, and PDF report generation capabilities accessible to non-technical researchers.",
        styles['BodyText']
    ))
    
    # Figure 7 placeholder
    add_pdf_figure(story, 7, "Streamlit Application",
        "Screenshot of web interface", styles, image_folder)
    
    # Figure 8 placeholder
    add_pdf_figure(story, 8, "PDF Report Output",
        "Generated PDF report example", styles, image_folder)
    
    story.append(PageBreak())
    
    # ========== 4. DISCUSSION ==========
    
    story.append(Paragraph("4. Discussion", styles['SectionHeading']))
    
    story.append(Paragraph("4.1 Strengths", styles['SubsectionHeading']))
    add_pdf_paragraphs(story, DISCUSSION["strengths"], styles)
    
    story.append(Paragraph("4.2 Limitations", styles['SubsectionHeading']))
    add_pdf_paragraphs(story, DISCUSSION["limitations"], styles)
    
    story.append(PageBreak())
    
    # ========== 5. FUTURE WORK ==========
    
    story.append(Paragraph("5. Future Work", styles['SectionHeading']))
    add_pdf_paragraphs(story, FUTURE_WORK, styles)
    
    story.append(PageBreak())
    
    # ========== 6. CONCLUSION ==========
    
    story.append(Paragraph("6. Conclusion", styles['SectionHeading']))
    add_pdf_paragraphs(story, CONCLUSION, styles)
    
    # Summary table
    table_data = [
        ["Metric", "Value"],
        ["Top-1 Accuracy", "72.26%"],
        ["Top-5 Accuracy", "85.55%"],
        ["Species Covered", "54"],
        ["Real-world Tests", "2/3 (66.7%)"],
    ]
    add_pdf_table(story, table_data, "Table 7: Summary Metrics", styles)
    
    story.append(PageBreak())
    
    # ========== 7. REFERENCES ==========
    
    story.append(Paragraph("7. References", styles['SectionHeading']))
    
    for i, ref in enumerate(REFERENCES, 1):
        ref_style = ParagraphStyle(
            name=f'Ref{i}',
            parent=styles['Normal'],
            fontSize=10,
            leftIndent=36,
            firstLineIndent=-36,
            spaceAfter=6
        )
        story.append(Paragraph(f"[{i}] {ref}", ref_style))
    
    story.append(PageBreak())
    
    # ========== 8. APPENDIX ==========
    
    story.append(Paragraph("8. Appendix", styles['SectionHeading']))
    
    story.append(Paragraph("A. Species List (Sample - First 20)", styles['SubsectionHeading']))
    
    # Species table (first 20 only for PDF)
    species_table_data = [["#", "Scientific Name", "English Name", "Samples"]]
    for i, (sci, eng, samples) in enumerate(SPECIES_LIST[:20], 1):
        species_table_data.append([str(i), sci[:25], eng[:20], str(samples)])
    
    add_pdf_table(story, species_table_data, "Table A1: Species List (Sample)", styles)
    
    story.append(Paragraph("B. Software Stack", styles['SubsectionHeading']))
    
    table_data = [
        ["Component", "Technology"],
        ["Deep Learning", "PyTorch 2.10, timm"],
        ["Audio Processing", "librosa 0.10"],
        ["Web Interface", "Streamlit"],
        ["PDF Generation", "ReportLab"],
    ]
    add_pdf_table(story, table_data, "Table B1: Software", styles)
    
    # Build PDF
    doc.build(story)
    
    print(f"✅ PDF saved: {pdf_path}")
    return pdf_path


def add_pdf_paragraphs(story, text, styles):
    """Add paragraphs to PDF story with proper formatting"""
    paragraphs = text.strip().split('\n\n')
    for para in paragraphs:
        if para.strip():
            # Handle bullet points
            if para.strip().startswith('•'):
                lines = para.strip().split('\n')
                for line in lines:
                    if line.strip():
                        bullet_text = line.strip().lstrip('•').strip()
                        bullet_style = ParagraphStyle(
                            name=f'Bullet_{id(line)}',
                            parent=styles['Normal'],
                            fontSize=11,
                            leftIndent=36,
                            bulletIndent=18,
                            spaceAfter=6,
                            fontName='Times-Roman'
                        )
                        story.append(Paragraph(f"• {bullet_text}", bullet_style))
            else:
                # Regular paragraph with proper indentation
                para_style = ParagraphStyle(
                    name=f'BodyText_{id(para)}',
                    parent=styles['Normal'],
                    fontSize=11,
                    leading=16,
                    alignment=TA_JUSTIFY,
                    fontName='Times-Roman',
                    firstLineIndent=36,
                    spaceAfter=12,
                    leftIndent=0,
                    rightIndent=0
                )
                story.append(Paragraph(para.strip(), para_style))


def add_pdf_table(story, data, caption, styles):
    """Add professionally formatted table to PDF story"""
    
    if not data:
        return
    
    # Calculate column widths based on content
    num_cols = len(data[0])
    col_width = 450 / num_cols
    col_widths = [col_width] * num_cols
    
    # Create table
    table = Table(data, colWidths=col_widths)
    
    # Professional table styling
    table_style = TableStyle([
        # Header styling
        ('BACKGROUND', (0, 0), (-1, 0), colors.Color(0.2, 0.4, 0.6)),  # Dark blue
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
        ('VALIGN', (0, 0), (-1, 0), 'MIDDLE'),
        ('FONTNAME', (0, 0), (-1, 0), 'Times-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 10),
        
        # Body styling
        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
        ('TEXTCOLOR', (0, 1), (-1, -1), colors.black),
        ('ALIGN', (0, 1), (-1, -1), 'LEFT'),
        ('VALIGN', (0, 1), (-1, -1), 'MIDDLE'),
        ('FONTNAME', (0, 1), (-1, -1), 'Times-Roman'),
        ('FONTSIZE', (0, 1), (-1, -1), 9),
        
        # Cell padding for better readability
        ('LEFTPADDING', (0, 0), (-1, -1), 8),
        ('RIGHTPADDING', (0, 0), (-1, -1), 8),
        ('TOPPADDING', (0, 0), (-1, -1), 8),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
        
        # Grid lines for clarity
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ('LINEABOVE', (0, 0), (-1, 0), 1, colors.black),
        ('LINEBELOW', (0, 0), (-1, 0), 1, colors.black),
        ('LINEBELOW', (0, -1), (-1, -1), 1, colors.black),
        
        # Alternate row colors for better readability
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.Color(0.95, 0.95, 1.0)]),
    ])
    
    table.setStyle(table_style)
    
    # Add to story
    story.append(Spacer(1, 12))
    story.append(table)
    story.append(Spacer(1, 6))
    story.append(Paragraph(caption, styles['TableCaption']))


def add_pdf_figure(story, num, title, description, styles, image_folder):
    """
    Add figure to PDF - loads actual image if available, otherwise shows placeholder
    
    Expected image naming: Figure_1.png, Figure_2.jpg, etc.
    Store images in Report_images folder
    """
    
    image_found = False
    image_path = None
    
    # Try different image formats
    for ext in ['.png', '.jpg', '.jpeg', '.gif']:
        potential_path = image_folder / f"Figure_{num}{ext}"
        if potential_path.exists():
            image_path = str(potential_path)
            image_found = True
            break
    
    story.append(Spacer(1, 12))
    
    if image_found and image_path:
        # Load and display actual image
        try:
            img = Image(image_path, width=5*inch, height=3.75*inch)
            story.append(img)
        except Exception as e:
            # Fallback to placeholder if image loading fails
            placeholder_data = [[f"IMAGE: {title}\n\n[Image file could not be loaded]\n[{description}]\n\n"]]
            table = Table(placeholder_data, colWidths=[450])
            table.setStyle(TableStyle([
                ('BOX', (0, 0), (-1, -1), 1, colors.grey),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('FONTNAME', (0, 0), (-1, -1), 'Times-Roman'),
                ('FONTSIZE', (0, 0), (-1, -1), 10),
                ('BACKGROUND', (0, 0), (-1, -1), colors.Color(0.97, 0.97, 0.97)),
                ('TOPPADDING', (0, 0), (-1, -1), 20),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 20),
            ]))
            story.append(table)
    else:
        # Show placeholder for missing images
        placeholder_data = [[f"[IMAGE PLACEHOLDER]\n\nFigure {num}: {title}\n\n{description}\n\n"
                           f"Place image as: Report_images/Figure_{num}.png"]]
        table = Table(placeholder_data, colWidths=[450])
        table.setStyle(TableStyle([
            ('BOX', (0, 0), (-1, -1), 2, colors.orange),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('FONTNAME', (0, 0), (-1, -1), 'Times-Roman'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BACKGROUND', (0, 0), (-1, -1), colors.Color(1.0, 0.97, 0.9)),
            ('TOPPADDING', (0, 0), (-1, -1), 20),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 20),
        ]))
        story.append(table)
    
    story.append(Spacer(1, 6))
    story.append(Paragraph(f"<b>Figure {num}:</b> {title}", styles['FigureCaption']))


# ============================================================
# MAIN FUNCTION
# ============================================================

def main():
    """Generate both DOCX and PDF reports"""
    
    print("=" * 60)
    print("📄 TECHNICAL REPORT GENERATOR")
    print("=" * 60)
    print()
    
    # Create output directory
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    print(f"📁 Output directory: {OUTPUT_DIR}")
    print()
    
    # Generate DOCX
    try:
        docx_path = create_docx_report()
    except Exception as e:
        print(f"❌ DOCX generation failed: {e}")
        docx_path = None
    
    print()
    
    # Generate PDF
    try:
        pdf_path = create_pdf_report()
    except Exception as e:
        print(f"❌ PDF generation failed: {e}")
        pdf_path = None
    
    # Summary
    print()
    print("=" * 60)
    print("📊 GENERATION COMPLETE")
    print("=" * 60)
    print()
    
    if docx_path:
        print(f"✅ DOCX: {docx_path}")
    if pdf_path:
        print(f"✅ PDF:  {pdf_path}")
    
    print()
    print("📝 Next steps:")
    print("   1. Open the DOCX file in Microsoft Word")
    print("   2. Replace figure placeholders with actual images")
    print("   3. Update author name and institution")
    print("   4. Generate table of contents")
    print("   5. Final formatting and review")
    print()


if __name__ == "__main__":
    main()